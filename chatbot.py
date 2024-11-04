from flask import Flask, render_template, request, jsonify, session, redirect, url_for
import os

from flask.cli import load_dotenv
from openai import embeddings
import torch
from bs4 import BeautifulSoup
import requests
from transformers import CLIPProcessor, CLIPModel
from PIL import Image
import chromadb
from langchain.prompts import ChatPromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import WebBaseLoader, PyPDFLoader
from langchain_chroma import Chroma
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from docx import Document as DocxDocument
from langchain_core.documents import Document
import concurrent.futures
import numpy as np
import logging
from langchain_core.runnables import RunnableLambda
import re
import threading
import spacy
from fuzzywuzzy import fuzz
from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(level=logging.WARNING)
logging.getLogger("httpx").setLevel(logging.WARNING)

app = Flask(__name__)
app.secret_key = 'your_secret_key'
app.config['UPLOAD_FOLDER'] = 'uploads'

device = "cuda" if torch.cuda.is_available() else "cpu"
clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32").to(device)
clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")

CHROMA_DB_DIR = "chromadb_storage"
client = None


def get_chroma_client():
    """Initialisiert und gibt den Chroma-Client zurück."""
    global client
    if client is None:
        client = chromadb.Client(chromadb.config.Settings(
            persist_directory=CHROMA_DB_DIR
        ))
    return client


chroma_client = get_chroma_client()
print("Starte mit Initialisierung für die Bilderkennung")
clip_collection = client.get_or_create_collection("clip_collection")

print("Starte mit Initialisierung für die Texterekknung")
text_collection = client.get_or_create_collection("waste_info")

# Wichtig für routing zum Ladebildschirm
setup_done = False
nlp = spacy.load("de_core_news_sm")


def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)


def setup_documents():
    global setup_done, vectorstore, retriever, llm, prompt_template
    print("Starte Setup")

    if not os.path.exists(CHROMA_DB_DIR):
        os.makedirs(CHROMA_DB_DIR)

    # Überprüfen, ob der Chroma-Datenbankordner leer ist (erstmaliger Setup)
    if not os.listdir(CHROMA_DB_DIR):
        print("Erstmaliges Setup: Dokumente werden geladen und verarbeitet...")

        docs = []

        # Word-Dokumente laden
        print("Beginne mit Word-Dokumenten")
        word_folder = os.path.join("data", "word_docs")
        word_file_paths = [os.path.join(word_folder, f) for f in os.listdir(word_folder) if f.endswith('.docx')]

        for word_path in word_file_paths:
            if not os.path.exists(word_path):
                print(f"Datei nicht gefunden: {word_path}")
                continue
            word_text = extract_text_from_word(word_path)
            if word_text is None:
                print(f"Kein Inhalt in Dokument: {word_path}")
                continue
            doc_obj = Document(page_content=word_text, metadata={"source": word_path})
            docs.append(doc_obj)
            print(f"Dokument hinzugefügt: {word_path}")

        # PDF-Dokumente laden
        print("Beginne mit PDF-Dokumenten")
        pdf_folder = os.path.join("data", "pdfs")
        pdf_file_paths = [os.path.join(pdf_folder, f) for f in os.listdir(pdf_folder) if f.endswith('.pdf')]

        for pdf_path in pdf_file_paths:
            if not os.path.exists(pdf_path):
                print(f"PDF-Datei nicht gefunden: {pdf_path}")
                continue
            pdf_loader = PyPDFLoader(pdf_path)
            pdf_docs = pdf_loader.load()  # PDF in Seiten aufspalten und laden

            for pdf_doc in pdf_docs:
                doc_obj = Document(page_content=pdf_doc.page_content, metadata={"source": pdf_path})
                docs.append(doc_obj)
        print(f"PDF-Dokument hinzugefügt: {pdf_path}")

        # Pfand-Website
        print("Beginne mit Websites. Dies kann einen Moment dauern.")
        pfandpflicht_text = scrape_pfandpflicht_website()
        if pfandpflicht_text:
            doc_obj = Document(page_content=pfandpflicht_text, metadata={"source": "Pfandpflicht Webseite"})
            docs.append(doc_obj)

        #Abfall-Abc
        abc_text = scrape_fes_abfall_abc()
        if abc_text:
            doc_obj = Document(page_content=abc_text, metadata={"source": "Abfall ABC"})
            docs.append(doc_obj)

        results = text_collection.get()

        for doc, metadata in zip(results['documents'], results['metadatas']):
            doc_obj = Document(
                page_content=doc,
                metadata={
                    'search_text': metadata['search_text'],
                    'synonyms': metadata['synonyms'],
                    'title': metadata['title'],
                    'subtitle': metadata['subtitle']
                }
            )
            docs.append(doc_obj)

        # Text in Chunks aufteilen
        print("Starte Chunking")
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        splits = text_splitter.split_documents(docs)

        print("Speichern in Vektordatenbank")
        vectorstore = Chroma(persist_directory=CHROMA_DB_DIR, embedding_function=get_embedding_function())
        vectorstore.add_documents(splits)
        print(f"Anzahl der gespeicherten Dokumente im Vektorstore: {len(splits)}")

    else:
        print("Persistente Dokumente bereits geladen, überspringe diesen Schritt.")
        # Lade den bestehenden Vektorstore für persistente Dokumente
        vectorstore = Chroma(persist_directory=CHROMA_DB_DIR, embedding_function=get_embedding_function())

    # Webseiten-Inhalte bei jedem Start neu laden (frisst performance, hält aber insb. Öffnungszeiten aktuell)
    print("Scrape Websites erneut. Dies kann einen Moment dauern.")
    scrape_and_add_websites()

    # Setze den Retriever und das Sprachmodell
    print("Feinschliff")
    retriever = vectorstore.as_retriever()
    llm = ChatOpenAI(model_name="gpt-4o-mini-2024-07-18", temperature=0)

    # Prompt Definieren
    PROMPT_TEMPLATE = """
    You are a friendly assistant that provides detailed but concise answers.
    Always respond in the same language as the user. If the user's question is in German, respond in German. If the user's question is in English, respond in English.
 
    If the user asks a question about waste management, provide a clear and concise answer based on the following context:
 
    {context}
 
    If the user greets you (e.g., 'hello', 'hi', 'hallo'), respond politely and ask how you can assist them with waste management.
 
    If the user's question is unclear or not related to waste management, politely ask for clarification or suggest that they provide more details.
 
    ---
 
    Please answer the following question based on the above context: {question}
 
    If you cannot answer the question based on the context, kindly ask the user for more details or offer alternative resources. Always respond in the language of the user.
    """

    global prompt_template
    prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)

    setup_done = True
    print("Setup abgeschlossen.")
    print("Chatbot ist nun Einsatzbereit.")


def scrape_and_add_websites():
    # Neue gescrapten Webseiteninhalte
    new_docs = []

    # Scrape die FES Abfall-ABC Seite
    scrape_fes_abfall_abc_content = scrape_fes_abfall_abc()
    if scrape_fes_abfall_abc_content:
        new_docs.append(Document(page_content=scrape_fes_abfall_abc_content, metadata={"source": "FES Abfall-ABC"}))

    # Scrape die Pfandpflicht Webseite
    pfandpflicht_content = scrape_pfandpflicht_website()
    if pfandpflicht_content:
        new_docs.append(Document(page_content=pfandpflicht_content, metadata={"source": "Pfandpflicht Webseite"}))

    # Text in Chunks aufteilen, falls die gescrapten Inhalte groß sind
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    splits = text_splitter.split_documents(new_docs)

    # Füge die neuen gescrapten Inhalte zum bestehenden Vektorstore hinzu
    if splits:
        vectorstore.add_documents(splits)
        print(f"{len(splits)} gescrapten Dokumente wurden erfolgreich hinzugefügt.")
    else:
        print("Keine neuen gescrapten Inhalte gefunden.")


# Unnötig
def get_embedding_function():
    return OpenAIEmbeddings()


def normalize_embedding(embedding):
    norm = np.linalg.norm(embedding)
    return (embedding / norm).tolist() if norm > 0 else embedding.tolist()


# Parallele Verarbeitung von Embeddings
def calculate_embeddings_in_parallel(splits):
    embedding_function = get_embedding_function()

    def embed_document(doc):
        # Welches Modell?
        return normalize_embedding(embedding_function.embed_query(doc.page_content))

    # Unnötig. Durch Langchain wird bereits parallelisiert
    with concurrent.futures.ThreadPoolExecutor() as executor:
        embeddings = list(executor.map(embed_document, splits))

    return embeddings


def extract_text_from_word(docx_path):
    doc = DocxDocument(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)


def extract_text_from_pdf(pdf_path):
    try:
        pdf_loader = PyPDFLoader(pdf_path)
        pdf_docs = pdf_loader.load()  # PDF in Seiten aufspalten und laden

        # Text aus allen Seiten extrahieren und zusammenfügen
        full_text = []
        for page_doc in pdf_docs:
            full_text.append(page_doc.page_content)

        return "\n".join(full_text).strip()  # Den gesamten Text zusammenfügen
    except Exception as e:
        print(f"Fehler beim Extrahieren von Text aus PDF {pdf_path}: {e}")
        return None


def scrape_fes_abfall_abc():
    url = "https://www.fes-frankfurt.de/informatives-frankfurtplus/abfall-abc"
    response = requests.get(url)

    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')
        waste_entries = soup.find_all('div', class_='col-sm-4 entry')
        data = []
        scraped_content = ""

        for entry in waste_entries:
            data_searchtext = entry.get('data-searchtext', '').strip()
            data_synonyms = entry.get('data-synonyms', '').strip()
            h2_tag = entry.find('h2').get_text(strip=True) if entry.find('h2') else ''
            h3_tag = entry.find('h3').get_text(strip=True) if entry.find('h3') else ''
            p_tags = entry.find_all('p')
            p_texts = '\n'.join([p.get_text(strip=True) for p in p_tags if p.get_text(strip=True)])

            data_entry = {
                'search_text': data_searchtext,
                'synonyms': data_synonyms,
                'title': h2_tag,
                'subtitle': h3_tag,
                'disposal_info': p_texts or 'Keine Entsorgungsinformationen verfügbar'
            }

            document_content = f"Abfall: {data_entry['title']}\nSynonyme: {data_entry['synonyms']}\nEntsorgungsinformation: {data_entry['disposal_info']}"
            data_entry['document_content'] = document_content

            scraped_content += document_content + "\n"

            text_collection.add(
                documents=[data_entry['document_content']],
                metadatas=[{
                    'search_text': data_entry['search_text'],
                    'synonyms': data_entry['synonyms'],
                    'title': data_entry['title'],
                    'subtitle': data_entry['subtitle']
                }],
                ids=[data_entry['search_text']]
            )

            results = text_collection.get(ids=[data_entry['search_text']])

        print("Erfolgreich Daten gescraped und in die Chroma-Sammlung eingefügt.")
        return scraped_content
    else:
        print(f"Fehler beim Abrufen der Seite. Statuscode: {response.status_code}")
        return None


def scrape_pfandpflicht_website():
    url = "https://landwirtschaft.hessen.de/umwelt/abfall-und-recycling/pfandpflicht"
    response = requests.get(url)

    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')

        # Versuch, den Inhalt innerhalb von bestimmten Klassen oder IDs zu finden
        main_content = soup.find('div', {'class': 'main-content'})  # Beispiel für eine mögliche Klasse
        if main_content is None:
            # Wenn keine 'main-content' gefunden wird, probiere eine allgemeinere Auswahl
            main_content = soup.find('main', {'role': 'main'})

        if main_content:
            # Extrahiere den Text
            text_content = main_content.get_text(separator='\n', strip=True)
            return text_content
        else:
            print("Hauptinhalt konnte nicht gefunden werden.")
            return None
    else:
        print(f"Fehler beim Abrufen der Seite. Statuscode: {response.status_code}")
        return None


def get_chat_response(question):
    try:
        # Abfrage nach Abfallkalender (z.B. wenn Frage nach Abholungstermin gestellt wird)
        if "wann" in question.lower() and ("müll" in question.lower() or "abholung" in question.lower()):
            return 'Hier kannst du den Abholkalender für Frankfurt einsehen: <a href="https://www.fes-frankfurt.de/services/abfallkalender" target="_blank">Abfallkalender FES Frankfurt</a>'

        # Abfrage relevanter Dokumente vom Retriever basierend auf der Frage
        retrieved_docs = retriever.get_relevant_documents(question)

        # Debugging: Zeige die abgerufenen Dokumente an
        print(f"Abgerufene Dokumente: {retrieved_docs}")
        print("CHECKPOINT")

        if not retrieved_docs:
            return "Keine relevanten Informationen gefunden."

        relevant_contexts = []

        for doc in retrieved_docs:
            if isinstance(doc, Document):
                search_text = doc.metadata.get("search_text", None)
                synonyms = doc.metadata.get("synonyms", [])

                # Stelle sicher, dass waste_item und synonyms überprüfbare Strings sind
                if isinstance(search_text, str) and search_text.strip():
                    valid_synonyms = [syn.strip() for syn in synonyms if isinstance(syn, str) and syn.strip()]

                    # Kombiniere waste_item und die validierten Synonyme in terms_to_check
                    terms_to_check = [search_text] + valid_synonyms

                    if is_relevant_fuzzy(question, terms_to_check):
                        relevant_contexts.append(doc.page_content)
                        print(f"Gefundener relevanter Kontext: {doc.page_content}")
                else:
                    # Falls keine Metadaten vorhanden sind, füge das Dokument hinzu
                    relevant_contexts.append(doc.page_content)
            else:
                logging.error(f"Unexpected document type: {type(doc)}. Skipping this document.")

        # Falls keine passenden Informationen gefunden wurden, gib eine entsprechende Nachricht zurück
        if not relevant_contexts:
            return "Es wurden keine spezifischen Informationen zu deiner Anfrage gefunden. Bitte gib mehr Details an oder frage nach anderen Gegenständen."

        # Sortiere und füge nur die wichtigsten Kontexte zusammen
        relevant_contexts = rank_by_relevance(relevant_contexts, question)
        context = "\n\n".join(relevant_contexts[:5])  # Verwende nur die 5 relevantesten Kontexte

        print("Zusammengeführter Kontext:", context)

        # Korrektes Übergabeformat für die RAG-Kette
        rag_chain = (
                {"context": RunnablePassthrough(), "question": RunnablePassthrough()}
                | prompt_template
                | llm
                | StrOutputParser()
        )

        # Antwort
        result = rag_chain.invoke({"context": context, "question": question})

        # Falls keine klare Antwort generiert wird, frage nach mehr Details
        if not result or "I'm not sure" in result:
            return f"Ich habe einige Informationen gefunden: {context}. Könntest du deine Frage näher erläutern?"

        return format_bold_line(result) if isinstance(result, str) else str(result)

    except Exception as e:
        logging.error(f"Error generating response: {e}")
        return "Ein Fehler ist aufgetreten. Bitte versuche es später noch einmal."


def format_bold_line(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)  # Fett
    return text.replace('\n', '<br>')  # Zeilenumbruch


def lemmatize_text(text):
    try:
        logging.debug(f"Text zur Lemmatisierung: {text}")
        doc = nlp(text)
        lemmas = [token.lemma_ for token in doc]
        logging.debug(f"Lemmatisierte Wörter: {lemmas}")
        return lemmas
    except Exception as e:
        logging.error(f"Fehler bei der Lemmatisierung: {str(e)}")
        return text.split()  # Fallback


def is_relevant_fuzzy(question, terms_to_check, threshold=70):
    question = question.lower()  # Setze die Frage auf Kleinbuchstaben, um Groß-/Kleinschreibung zu ignorieren
    for term in terms_to_check:
        term = term.lower()  # Setze auch die Suchbegriffe auf Kleinbuchstaben
        # Berechne die Ähnlichkeit zwischen der Frage und dem Suchbegriff
        similarity = fuzz.partial_ratio(question, term)
        print(f"Ähnlichkeit zwischen '{question}' und '{term}': {similarity}")

        if similarity >= threshold:
            print(f"Relevanter Begriff gefunden: {term} (Ähnlichkeit: {similarity}%)")
            return True

    print("Kein relevanter Begriff gefunden (Fuzzy Matching).")
    return False


def rank_by_relevance(relevant_contexts, question):
    def count_matches(context):
        context_words = set(context.lower().split())
        question_words = set(question.lower().split())
        return len(context_words & question_words)  # Anzahl der gemeinsamen Wörter

    return sorted(relevant_contexts, key=count_matches, reverse=True)


@app.route("/upload", methods=["GET", "POST"])
def upload_image():
    if 'image' not in request.files:
        return jsonify({"error": "No file part"})

    file = request.files['image']
    if file.filename == '':
        return jsonify({"error": "No selected file"})

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(file_path)

    recognized_label = get_clip_label(file_path)

    refined_label = f"Entsorgung von {recognized_label}"
    context_label = f"Wie entsorge ich {recognized_label}? Informationen zur Abfallentsorgung in Frankfurt."
    print("##############")
    test = get_chat_response("Wie entsroge ich " + recognized_label)
    print(test)
    print("##############")

    clip_embedding = get_clip_embedding(context_label)
    normalized_clip_embedding = normalize_embedding(clip_embedding)

    # Retrieve relevant documents from clip_collection
    retrieved_docs = clip_collection.query(query_embeddings=[normalized_clip_embedding], n_results=10)
    print(retrieved_docs)
    document_contents = [doc for doc_list in retrieved_docs.get("documents", []) for doc in doc_list]
    print(document_contents)

    dynamic_contexts = []

    # Dynamische Kontexte basierend auf Metadaten generieren, falls vorhanden
    if "metadatas" in retrieved_docs:
        for metadata in retrieved_docs['metadatas']:
            dynamic_contexts += generate_dynamic_contexts_from_metadata(metadata)

    # Relevante Kontexte bewerten und sortieren, falls dynamische Kontexte vorhanden sind
    if dynamic_contexts:
        relevant_contexts = rank_by_relevance(dynamic_contexts, f"Wie entsorge ich {recognized_label}?")
    else:
        relevant_contexts = []

    print(recognized_label)

    question = "Wie entsorge ich " + recognized_label + "?"
    response = get_chat_response(question)
    print(response)

    return jsonify({
        "recognized_label": recognized_label,
        "retrieved_docs": response,
        "relevant_contexts": relevant_contexts  # Füge relevante Kontexte zur Antwort hinzu
    })


def generate_dynamic_contexts_from_metadata(metadata):
    fields = ['search_text', 'synonyms', 'title', 'subtitle']
    dynamic_contexts = []

    for field in fields:
        if field in metadata and metadata[field]:
            context_value = metadata[field]
            dynamic_contexts.append(f"Wie entsorge ich {context_value}?")
            dynamic_contexts.append(f"Was ist {context_value}?")
            dynamic_contexts.append(f"Wie benutze ich {context_value}?")

    return dynamic_contexts


def process_metadata_with_fallback(results):
    docs = []

    # Feste Metadatenfelder, die wir bevorzugen
    hardcoded_fields = ['search_text', 'synonyms', 'title', 'subtitle']

    for doc, metadata in zip(results['documents'], results['metadatas']):
        extracted_metadata = {}

        # Prüfe hartcodierte Felder und füge sie hinzu, wenn sie vorhanden sind
        for field in hardcoded_fields:
            if field in metadata and metadata[field]:
                extracted_metadata[field] = metadata[field]

        # Füge das Dokument zusammen mit den extrahierten Metadaten hinzu
        doc_obj = Document(page_content=doc, metadata=extracted_metadata)
        docs.append(doc_obj)

    return docs


def get_clip_embedding(text):
    truncated_text = truncate_text(text, max_length=77)
    inputs = clip_processor(text=[truncated_text], return_tensors="pt", padding=True, truncation=True,
                            max_length=77).to(device)
    with torch.no_grad():
        text_features = clip_model.get_text_features(input_ids=inputs['input_ids'],
                                                     attention_mask=inputs['attention_mask']).cpu().numpy()
    return text_features.squeeze().tolist()


def truncate_text(text, max_length=77):
    tokens = text.split()
    return ' '.join(tokens[:max_length])


def get_clip_label(image_path):
    try:
        image = Image.open(image_path)
        image = image.resize((224, 224))
    except Exception as e:
        print(f"Fehler beim Laden des Bildes: {e}")
        return None

    labels = [
        # Biomüll
        "Apfelschale", "Bananenschale", "Kaffeesatz", "Teebeutel", "Obst", "Gemüse", "Eierschalen", "Gartenabfälle",
        "Laub", "Grünabfall",
        "Altes Brot", "Knochen", "Papiertücher", "Zimmerpflanzen",

        # Elektronikschrott
        "Computer", "Laptop", "Smartphone", "Tablet", "Fernseher", "Drucker", "Kabel", "Ladegerät", "Kopfhörer",
        "Lautsprecher",
        "Mikrowelle", "Kühlschrank", "Waschmaschine", "Spülmaschine", "Toaster", "Staubsauger", "Bügeleisen",

        # Glas
        "Glasflasche", "Weinflasche", "Bierflasche", "Marmeladenglas", "Konservenglas", "Parfumflasche", "Vase",
        "Trinkglas", "Fensterglas", "Spiegelglas",

        # Plastik
        "Plastikflasche", "Plastiktüte", "Plastikbecher", "Joghurtbecher", "Shampooflasche", "Plastikverpackung",
        "Plastikspielzeug",
        "Kunststofffolie", "Frischhaltefolie", "Verpackungsmaterial", "Plastikbesteck", "Einwegplastikgeschirr",

        # Metall
        "Alufolie", "Konservendose", "Getränkedose", "Metallkanister", "Schrauben", "Nägel", "Metallfolie", "Besteck",
        "Kochtopf",
        "Pfanne", "Fahrrad", "Metallspielzeug",

        # Papier und Pappe
        "Papier", "Zeitung", "Karton", "Pappe", "Briefumschlag", "Katalog", "Buch", "Schreibpapier", "Notizbuch",
        "Toilettenpapierrolle",
        "Papiertüte", "Kartonverpackung",

        # Restmüll
        "Windeln", "Keramik", "Tasse", "Keramikteller", "Spiegel", "Glühbirne", "Zigarettenkippen", "Staubsaugerbeutel",
        "Kaugummi", "Einwegmaske", "Verbandmaterial", "Zahnbürste",

        # Sondermüll
        "Batterie", "Akkumulator", "Energiesparlampe", "Leuchtstoffröhre", "Farbe", "Lack", "Altöl", "Chemikalien",
        "Putzmittel", "Insektenspray", "Spraydose", "Feuerlöscher",

        # Sperrmüll
        "Holztisch", "Trampolin", "Sofa", "Bett", "Matratze", "Schrank", "Stuhl", "Teppich", "Regal", "Fahrrad",
        "Autoreifen", "Kinderwagen",

        # Textilien
        "Kleidung", "Hose", "Shirt", "Jacke", "Pullover", "Schuhe", "Schal", "Handschuhe", "Bettwäsche", "Kissen",
        "Decke",
        "Vorhang", "Teppich", "Gürtel", "Tasche",

        # Gelber Sack (Verpackungen)
        "Tetrapak", "Verpackungsfolie", "Plastikverpackung", "Chips-Tüte", "Joghurtbecherdeckel", "Aluminiumverpackung",
        "Snackverpackung",
        "Milchverpackung", "Getränkekarton", "Verpackungsstyropor", "Plastikschale", "Verpackungsband"
    ]
    inputs = clip_processor(images=image, return_tensors="pt").to(device)
    with torch.no_grad():
        image_features = clip_model.get_image_features(pixel_values=inputs['pixel_values'])

    text_inputs = clip_processor(text=labels, return_tensors="pt", padding=True).to(device)
    with torch.no_grad():
        text_features = clip_model.get_text_features(input_ids=text_inputs['input_ids'],
                                                     attention_mask=text_inputs['attention_mask'])

    logits_per_image = torch.matmul(image_features, text_features.t())
    probs = logits_per_image.softmax(dim=-1)

    predicted_label_index = probs.argmax().item()
    return labels[predicted_label_index]


@app.route("/")
def index():
    # Routing ladebildschrim
    return render_template('loading.html')


@app.route("/chat")
def chat_redirect():
    if setup_done:
        return render_template('chat.html')
    else:
        return redirect(url_for('index'))


@app.route("/get", methods=["GET", "POST"])
def chat():
    try:
        user_question = request.form["msg"]

        response = get_chat_response(user_question)

        if 'context' not in session:
            session['context'] = ""
        session['context'] += f"User: {user_question}\nBot: {response}\n"

        print(f"Backend-Antwort: {response}")  # Debugging

        if isinstance(response, str):
            return jsonify({"text": response})
        else:
            return jsonify(response)

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": f"An error occurred while processing your request: {str(e)}"})


setup_thread = threading.Thread(target=setup_documents)
setup_thread.start()

if __name__ == '__main__':
    app.run(debug=True)
